# app/utils/resume_parser.py
#
# Raw text extraction from uploaded resumes (PDF / DOCX / TXT).
# Extraction is local and cheap; semantic structuring is done by the LLM
# in resume_service (same split as audio: local signal extraction + LLM).

import os


class ResumeParseError(Exception):
    pass


def extract_resume_text(path: str, filename: str | None = None) -> str:
    """Extract plain text from a resume file based on its extension."""
    name = (filename or path).lower()

    if name.endswith(".pdf"):
        text = _extract_pdf(path)
    elif name.endswith(".docx"):
        text = _extract_docx(path)
    elif name.endswith(".txt"):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
    else:
        ext = os.path.splitext(name)[1] or "unknown"
        raise ResumeParseError(
            f"Unsupported resume format '{ext}'. Upload a PDF, DOCX or TXT file."
        )

    text = "\n".join(line.strip() for line in text.splitlines())
    text = text.strip()
    if len(text) < 100:
        raise ResumeParseError(
            "Could not extract readable text from the resume. "
            "The file may be scanned images or empty."
        )
    return text


def _extract_pdf(path: str) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ResumeParseError(
            "pypdf not installed. Run: pip install pypdf"
        )
    reader = PdfReader(path)
    pages = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(pages)


def _extract_docx(path: str) -> str:
    try:
        import docx
    except ImportError:
        raise ResumeParseError(
            "python-docx not installed. Run: pip install python-docx"
        )
    document = docx.Document(path)
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)
